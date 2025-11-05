"""
python-docx==0.8.11
transformers==4.30.0
torch==2.0.0
torchvision==0.15.0
sentence-transformers==2.2.2
Pillow==9.5.0
pytesseract==0.3.10
opencv-python==4.7.0.72
numpy==1.24.0
chromadb==0.3.21
timm=1.0.12
"""

import os
import json
import base64
from datetime import datetime
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional
from PIL import Image
import pytesseract
from docx import Document
import torch
from transformers import (
    AutoTokenizer, AutoModel, 
    CLIPProcessor, CLIPModel,
    TableTransformerForObjectDetection
)
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings

@dataclass
class ContentBlock:
	"""内容块数据结构"""
	block_id: str
	content_type: str # "text","table","image"
	raw_data: Any
	processed_data: Optional[str]=None
	metadata: Dict[str,Any]=field(default_factory=dict)
	embeddings: Dict[str, List[float]] = field(default_factory=dict)
	position_info: Dict[str, Any] = field(default_factory=dict)

@dataclass
class DocumentMetadata:
	"""文档元数据"""
	doc_id:str
	filename:str
	processed_time:str
	total_blocks:int=0
	content_blocks:List[ContentBlock]=field(default_factory=list)

class MultiModalEmbeddingConfig:
	"""多模态嵌入配置"""
	def __init__(self):
		# 文本嵌入模型
		self.text_model_name="sentence-transformers/all-mpnet-base-v2"
		# 图像嵌入模型
		self.image_model_name="openai/clip-vit-base-patch32"
		# 表格检测模型
		self.table_detection_name="microsoft/table-transformer-detection"

		# OCR配置
		self.ocr_lang="eng+chi_sim" # 英语+简体中文

		# 嵌入维度
		self.text_embedding_dim=1024
		self.image_embedding_dim=512

		# 处理参数
		self.max_text_length=2000
		self.image_size=(224,224)

class DocumentProcessor:
	"""文档处理"""
	def __init__(self,config:MultiModalEmbeddingConfig):
		self.config=config
		self._load_models()

	def _load_models(self):
		"""加载所有需要的模型"""
		print("加载模型中...")
		
		# 文本嵌入模型
		self.text_model=SentenceTransformer(self.config.text_model_name)

		# 图像模型
		self.clip_model=CLIPModel.from_pretrained(self.config.image_model_name)
		self.clip_processor=CLIPProcessor.from_pretrained(self.config.image_model_name)

		# 表格检测模型
		self.table_detector=TableTransformerForObjectDetection.from_pretrained(self.config.table_detection_name)

		# 文本分词器
		self.tokenizer=AutoTokenizer.from_pretrained(self.config.text_model_name)

		print("模型加载完成")

class ContentExtractor:
	"""文件内容提取器"""
	def __init__(self,processor:DocumentProcessor):
		self.processor=processor

	def extract_from_docx(self,file_path:str)->DocumentMetadata:
		"""从docx文档中提取内容"""
		print(f"从{file_path}中提取内容")

		doc_id=self._generate_doc_id(file_path)
		# 创建metadata对象
		metadata=DocumentMetadata(
			doc_id=doc_id,
			filename=os.path.basename(file_path),
			processed_time=datetime.now().isoformat()
		)

		# 读取文档
		doc=Document(file_path)
		# 提取所有内容块
		content_blocks=self._extract_content_blocks(doc,file_path)
		metadata.content_blocks=content_blocks
		metadata.total_blocks=len(content_blocks)

		return metadata
	
	def _generate_doc_id(self,file_path:str)->str:
		"""生成文档ID"""
		filename=os.path.basename(file_path)
		timestamp=datetime.now().strftime("%Y%m%d_%H%M%S")
		return f"{filename}_{timestamp}"

	def _extract_content_blocks(self,doc:Document,file_path:str)->List[ContentBlock]:
		"""提取所有内容块"""
		blocks=[]
		block_index=0

		# 提取段落文本
		for para_idx,paragraph in enumerate(doc.paragraphs):
			if paragraph.text.strip():
				block=ContentBlock(
					block_id=f"text_{block_index:05d}",
					content_type="text",
					raw_data=paragraph.text.strip(),
					position_info={"paragraph_index":para_idx}
				)
				blocks.append(block)
				block_index+=1
		# 提取表格
		for table_idx,table in enumerate(doc.tables):
			table_data=self._extract_table_data(table)
			block=ContentBlock(
				block_id=f"table_{block_index:05d}",
				content_type="table",
				raw_data=table_data,
				metadata={"row_count":len(table_data),"col_count":len(table_data[0]) if table_data else 0},
				position_info={"table_index":table_idx}
			)
			blocks.append(block)
			block_index+=1
		# 提取图片
		images=self._extract_images_from_docx(file_path)
		for img_idx,(img_path,img_info) in enumerate(images):
			block=ContentBlock(
				block_id=f"image_{block_index:05d}",
				content_type="image",
				raw_data=img_path,
				metadata=img_info,
				position_info={"image_index":img_idx}
			)
			blocks.append(block)
			block_index+=1

		return blocks

	def _extract_table_data(self,table)->List[List[str]]:
		"""提取表格内容"""
		table_data=[]
		for row in table.rows:
			row_data=[cell.text.strip() for cell in row.cells]
			table_data.append(row_data)
		return table_data

	def _extract_images_from_docx(self,file_path:str)->List[tuple]:
		"""提取图片"""
		import zipfile
		import tempfile
		import shutil

		images=[]
		temp_dir=tempfile.mkdtemp()

		try:
			# 加压docx文件
			with zipfile.ZipFile(file_path,"r") as docx_zip:
				# 提取所有图片文件
				image_files=[]
				for f in docx_zip.namelist():
					if f.startswith("word/media") and f.split(".")[-1].lower() in ["png","jpg","jpeg","gif","bmp"]:
						image_files.append(f)
				for image_file in image_files:
					# 提取图片到临时目录
					extracted_path=os.path.join(temp_dir,os.path.basename(image_file))
					with docx_zip.open(image_file) as source,open(extracted_path,"wb") as target:
						target.write(source.read())
					image_info=self._process_image_info(extracted_path)
					images.append((extracted_path,image_info))
		except Exception as e:
			print(f"提取图像时出错:{e}")
		finally:
			# shutil.rmtree(temp_dir)
			pass
		return images

	def _process_image_info(self,image_path:str)->Dict[str,Any]:
		"""处理图片信息"""
		try:
			with Image.open(image_path) as img:
				# OCR文字识别

				# TODO 这里还可以使用大预言模型识别，如qwen-vl
				ocr_text=pytesseract.image_to_string(img,lang=self.processor.config.ocr_lang)

				#图片信息
				img_info={
					"width":img.width,
					"height":img.height,
					"format":img.format,
					"mode":img.mode,
					"ocr_text":ocr_text.strip(),
					"file_size":os.path.getsize(image_path)
				}

				return img_info
		except Exception as e:
			print(f"处理{image_path}图像时出错:{e}")
			return {"error":str(e)}

class EmbeddingGenerator:
	"""多模态嵌入生成器"""
	def __init__(self,processor:DocumentProcessor):
		self.processor=processor

	def generate_embedding(self,document_metadata:DocumentMetadata)->DocumentMetadata:
		"""为文档生成多模态嵌入"""
		print("生成多模态嵌入...")
		for block in document_metadata.content_blocks:
			try:
				if block.content_type=="text":
					self._generate_text_embedding(block)
				elif block.content_type=="table":
					self._generate_table_embedding(block)
				elif block.content_type=="image":
					self._generate_image_embedding(block)
				# 生成组合嵌入
				self._generate_combined_embedding(block)
			except Exception as e:
				print(f"生成块嵌入时出错,块编号：{block.block_id}:{e}")

		return document_metadata
	def _generate_text_embedding(self,block:ContentBlock):
		"""生成文本嵌入"""

		#预处理文本
		processed_text=self._preprocess_text(block.raw_data)
		block.processed_data=processed_text

		embedding=self.processor.text_model.encode([processed_text])[0]
		block.embeddings["text"]=embedding.tolist()

	def _generate_table_embedding(self,block:ContentBlock):
		"""生成表格嵌入"""
		# 将表格转为可描述文本
		table_description=self._table_to_desctiption(block.raw_data)
		block.processed_data=table_description

		# 使用文本模型生成嵌入
		embedding=self.processor.text_model.encode([table_description])[0]
		block.embeddings["table"]=embedding.tolist()

	def _generate_image_embedding(self,block:ContentBlock):
		"""生成图片嵌入"""
		try:
			# 加载图片
			image=Image.open(block.raw_data)

			# 使用clip模型生成图片嵌入
			inputs=self.processor.clip_processor(images=image,return_tensors="pt",padding=True)

			with torch.no_grad():
				image_features=self.processor.clip_model.get_image_features(**inputs)
				image_embedding=image_features[0].cpu().numpy()

			block.embeddings["image"]=image_embedding.tolist()

			# 生成图片描述
			ocr_text=block.metadata.get("ocr_text","")
			if ocr_text:
				text_embedding=self.processor.text_model.encode([ocr_text])[0]
				block.embeddings["image_text"]=text_embedding.tolist()

		except Exception as e:
			print(f"生成图片{block.raw_data}嵌入时错误:{e}")

	def _generate_combined_embedding(self,block:ContentBlock):
		"""生成组合嵌入"""
		combined_embedding=[]

		# 收集所有可用嵌入
		embeddings_to_combine=[]

		if "text" in block.embeddings:
			embeddings_to_combine.append(block.embeddings["text"])
		elif "table" in block.embeddings:
			embeddings_to_combine.append(block.embeddings["table"])
		elif "image_text" in block.embeddings:
			embeddings_to_combine.append(block.embeddings["image_text"])
		elif "image" in block.embeddings:
			# 将图片嵌入投影到文本嵌入相似的维度
			image_embedding=self._project_image_embedding(block.embeddings["image"])
			embeddings_to_combine.append(image_embedding)

		# 简单平均融合
		if embeddings_to_combine:
			combined_embedding=np.mean(embeddings_to_combine,axis=0).tolist()
			block.embeddings["combined"]=combined_embedding

	def _preprocess_text(self,text:str)->str:
		"""预处理文本"""

		# 简单的文本清理
		text=" ".join(text.split()) #移除多余的空格

		# 这里是写死了截断文本的长度，截断后可能会导致文本信息丢失！！！
		# 后续可以使用下面的智能截断的方法_smart_truncate，_split_text_into_chunks，_semantic_chunking
		if len(text)>self.processor.config.max_text_length:
			text=text[:self.processor.config.max_text_length]+'...'

		return text
	def _smart_truncate(self,text:str,max_length:int)->str:
		"""智能截断文本，尽量在句子边界处截断"""

		# 在句子边界处截断
		# if "。" in text[:max_length]:
		# 	# 找到最后一个句号
		# 	last_period=text[:max_length].rfind("。")
		# 	if last_period>max_length*0.8: # 确保截断点不要太靠前
		# 		return text[:last_period+1]+'...'

		# 在段落边界截断
		# if "\n" in text[:max_length]:
		# 	last_newline=text[:max_length].rfind("\n")
		# 	if last_newline>max_length*0.8:
		# 		return text[:last_newline].strip()+'...'

		# 在词语边界截断
		# if " " in text[:max_length]:
		# 	last_space=text[:max_length].rfind(" ")
		# 	if last_space>max_length*0.9:
		# 		return text[:last_space]+'...'

		# 硬截断
		return text[:max_length]+'...'
	def _split_text_into_chunks(self,text:str,overlap:int=50)->List[str]:
		"""将文本分割成重叠的块"""
		chunks=[]
		start=0
		max_len=self.processor.config.max_text_length

		while start<len(text):
			end=start+max_len
			if end<len(text):
				potential_endpoints=[
					text.rfind("。",start,end),
					text.rfind(".",start,end), #英文句号
					text.rfind("!",start,end),
					text.rfind("?",start,end),
					text.rfind("\n",start,end),
					text.rfind(" ",start,end),
				]

				valid_endpoints=[ep for ep in potential_endpoints if ep>start+max_len*0.6]

				if valid_endpoints:
					end=max(valid_endpoints)+1

			chunk=text[start:end].strip()
			
			if chunk:
				chunks.append(chunk)

			start=end-overlap if end-overlap>start else end

			if start>len(text):
				break	
		return chunks
	def _semantic_chunking(self,text:str)->List[str]:
		"""基于语义的文本分割"""
		import nltk
		from nltk.tokenize import sent_tokenize
		

		try:
			# 分句
			sentences=sent_tokenize(text)
		except:
			sentences=text.split(". ")

		chunks=[]
		current_chunk=""

		for sentence in sentences:
			# 如果当前块加上新句子不会超出最大长度
			if len(current_chunk)+len(sentence)<=self.processor.config.max_text_length:
				current_chunk+=sentence+" "
			else:
				# 保存当前块，并开始新块
				if current_chunk:
					chunks.append(current_chunk.strip())
				current_chunk=sentence+" "

		# 添加最后一个块
		if current_chunk:
			chunks.append(current_chunk.strip())

		return chunks
	
	def _table_to_description(self,table_data:List[List[str]])->str:
		"""将表格转化成描述性文本"""
		if not table_data:
			return "空表格"

		description="表格内容:\n"

		for i,row in enumerate(table_data):
			row_text="|".join([cell for cell in row if cell])
			description+=f"行{i+1}:{row_text}\n"

		return description

	def _project_image_embedding(self,image_embedding:List[float])->List[float]:
		"""将图像嵌入投影到文本嵌入相似的维度"""

		# 这里是简单的线性投影（实际应用中可能需要训练投影层）!如果不使用相似维度，则图片内容可以使用自己的维度保存到另外的库中，检索时同时检索文本和图片
		target_dim=self.processor.config.text_embedding_dim
		current_dim=len(image_embedding)

		if current_dim==target_dim:
			return image_embedding

		# 简单的重复或截断策略
		if current_dim<target_dim:
			# 重复嵌入直到达到目标维度
			repeat_times=target_dim//current_dim+1
			projected=(image_embedding*repeat_times)[:target_dim]
		else:
			# 截断到目标维度
			projected=image_embedding[:target_dim]

		return projected

if __name__=="__main__":

	mutiModalEmbeddingConfig=MultiModalEmbeddingConfig()
	
	documentProcessor=DocumentProcessor(mutiModalEmbeddingConfig)

	contentExtractor=ContentExtractor(documentProcessor)

	doc_metadata=contentExtractor.extract_from_docx("./source/迪士尼乐园 _园区地图、各主题区域介绍.docx")

	print(doc_metadata)
	for content_block in enumerate(doc_metadata.content_blocks):
		print("*"*100)
		print(content_block)
		